"""
============================================================================
 DOCUMENT PREPROCESSING PIPELINE  (Python)
============================================================================

 GOAL
 ----
 In a RAG (Retrieval-Augmented Generation) system, everything starts with
 RAW TEXT. But your knowledge lives in many formats: PDFs, Word docs,
 Markdown notes, plain text files...

 This script does ONE job: scan the ./docs folder and, for EVERY file it
 finds, pull out the plain text — regardless of the file's format.

 THE MENTAL MODEL
 ----------------
 It's an assembly line:

     docs/ folder ─▶ discover files ─▶ choose the right extractor
                                        (based on file extension)
                  ─▶ extract text   ─▶ collect results

 Each FORMAT needs its own extractor because the formats are stored very
 differently on disk:
   - .txt / .md  -> already plain text, just read it
   - .pdf        -> binary; a library walks its internal structure
   - .docx       -> secretly a ZIP of XML files; a library unpacks it
 We hide those differences behind a common idea: every extractor is a
 function that takes a file path and returns a string. That uniformity is
 the entire trick.
============================================================================
"""

# `pathlib` is the modern, object-oriented way to handle file paths in
# Python. A `Path` knows things about itself, like its `.suffix` (extension).
from pathlib import Path

# `dataclass` lets us define a small structured record with almost no
# boilerplate — just a labeled container for our results.
from dataclasses import dataclass

# Third-party libraries (installed via `pip install -r requirements.txt`).
# Each is a specialist that decodes one tricky format:
#   pypdf        -> reads PDF files
#   python-docx  -> reads .docx (Word) files (imported as `docx`)
from pypdf import PdfReader
from docx import Document


# ---------------------------------------------------------------------------
# A simple record describing one processed document.
# `@dataclass` auto-generates the __init__ etc. for us.
# ---------------------------------------------------------------------------
@dataclass
class ExtractedDocument:
    file_name: str   # e.g. "report.pdf"
    file_type: str   # the extension, e.g. ".pdf"
    text: str        # the extracted plain text
    char_count: int  # quick sanity check on how much we got


# The folder we scan. `__file__` is THIS script's path; `.parent` is the
# folder it lives in; then we point at the sibling "docs" folder.
DOCS_DIR = Path(__file__).parent / "docs"


# ===========================================================================
#  THE EXTRACTORS — one function per format.
#  They all share the SAME shape: take a Path, return a str. That sameness
#  is what lets the dispatcher treat them interchangeably.
# ===========================================================================

def extract_pdf(file_path: Path) -> str:
    """PDF — binary format. PdfReader parses its internal structure and
    lets us pull text page by page. We join the pages with newlines."""
    reader = PdfReader(str(file_path))
    # A PDF is a list of pages; extract_text() returns the text of each.
    # `or ""` guards against pages that have no extractable text (e.g.
    # scanned images), so we never try to join a None into the string.
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n".join(pages)


def extract_docx(file_path: Path) -> str:
    """DOCX — a Word document is actually a ZIP of XML files. python-docx
    unzips it and exposes the content as a list of paragraphs. We pull the
    text out of each paragraph and join them with newlines."""
    document = Document(str(file_path))
    paragraphs = [p.text for p in document.paragraphs]
    return "\n".join(paragraphs)


def extract_plain_text(file_path: Path) -> str:
    """Markdown (.md) and plain text (.txt) are ALREADY text on disk, so
    'extraction' is simply reading the file as UTF-8. We keep Markdown
    symbols (#, *, -) as-is; they're harmless for RAG and can even help
    preserve document structure."""
    # encoding="utf-8" decodes the bytes into a proper string.
    return file_path.read_text(encoding="utf-8")


# ===========================================================================
#  THE DISPATCHER — "given a file, which extractor do I use?"
#
#  We map each file extension to its extractor function. This is called a
#  "dispatch table". Adding a new format later = add ONE line to this dict.
#  Notice we store the FUNCTIONS themselves (no parentheses) as values.
# ===========================================================================
EXTRACTORS = {
    ".pdf": extract_pdf,
    ".docx": extract_docx,
    ".md": extract_plain_text,
    ".txt": extract_plain_text,
}


def extract_text(file_path: Path) -> str:
    """Look at the extension, find the matching extractor, and run it."""
    # `.suffix` gives ".PDF"; `.lower()` normalizes so ".PDF"/".Pdf"/".pdf"
    # all route to the same place.
    ext = file_path.suffix.lower()

    extractor = EXTRACTORS.get(ext)
    if extractor is None:
        # Raise loudly instead of silently returning "" — you want to KNOW
        # when a file was skipped because of an unsupported type.
        raise ValueError(f'Unsupported file type: "{ext}"')

    return extractor(file_path)


# ===========================================================================
#  THE PIPELINE — tie it all together.
# ===========================================================================
def process_all_documents() -> list[ExtractedDocument]:
    results: list[ExtractedDocument] = []

    # 1) DISCOVER + LOOP: iterdir() yields every entry in the folder.
    #    sorted() just makes the output order predictable.
    for file_path in sorted(DOCS_DIR.iterdir()):
        # Skip sub-folders and hidden files (like macOS's .DS_Store).
        if file_path.is_dir() or file_path.name.startswith("."):
            continue

        # 2) EXTRACT — wrapped in try/except so ONE bad file doesn't crash
        #    the whole run. Real folders are full of surprises; a production
        #    pipeline must degrade gracefully instead of exploding.
        try:
            text = extract_text(file_path)
            results.append(
                ExtractedDocument(
                    file_name=file_path.name,
                    file_type=file_path.suffix.lower(),
                    text=text,
                    char_count=len(text),
                )
            )
            print(f"✅ {file_path.name} — extracted {len(text)} characters")
        except Exception as err:
            # Report and move on to the next file.
            print(f"⚠️  Skipped {file_path.name}: {err}")

    return results


def main() -> None:
    print(f"📂 Scanning: {DOCS_DIR}\n")
    docs = process_all_documents()

    print(f"\n📊 Done. Processed {len(docs)} document(s).")

    # Show a short preview so you can SEE the text really came out.
    for doc in docs:
        # Collapse whitespace and trim to 120 chars for a clean one-liner.
        preview = " ".join(doc.text.split())[:120]
        print(f"\n--- {doc.file_name} ({doc.char_count} chars) ---")
        print(preview + (" ..." if doc.char_count > 120 else ""))

    # NEXT STEPS in a full RAG system (not done here):
    #   text -> CHUNK into small pieces -> EMBED each chunk -> STORE in a
    #   vector database. This script stops at clean text: the foundation.


# This guard means "only run main() when the file is executed directly"
# (i.e. `python main.py`), not when it's imported by another module.
if __name__ == "__main__":
    main()
